import json
from pathlib import Path

import docx
import pandas as pd
import pypdf
import tabula

from src.common.utils import Paths, _add_metadata_to_document


def _extract_tables_from_docx(table: docx.table.Table) -> pd.DataFrame:
    data = []
    keys = None
    for i, row in enumerate(table.rows):
        text = (cell.text for cell in row.cells)

        if i == 0:
            keys = tuple(text)
            continue
        row_data = dict(zip(keys, text))
        data.append(row_data)
    return pd.DataFrame(data).to_markdown()


def _extract_tables_from_pdf(file_path: Path) -> list[str]:
    tables = tabula.read_pdf(file_path, pages="all")
    return [table.to_markdown() for table in tables]


def extract_text_from_document(file_path: Path) -> str:
    if file_path.suffix == ".pdf":
        doc = pypdf.PdfReader(open(file_path, "rb"))
        text = "\n\n".join([page.extract_text() for page in doc.pages])
        tables = "\n\n".join(_extract_tables_from_pdf(file_path))

    elif file_path.suffix.startswith(".doc"):
        doc = docx.Document(file_path)
        text = "\n\n".join([para.text for para in doc.paragraphs])
        tables = "\n\n".join([_extract_tables_from_docx(table) for table in doc.tables])

    metadata = _add_metadata_to_document(file_path.stem)

    return (
        f"Dataset Title: {metadata['title']}\n\n"
        f"Dataset Notes: \n\n {metadata['notes']}\n\n"
        f"Document Text: \n\n {text}\n\n"
        f"Tables: \n\n {tables}"
    )


def add_metadata_to_document(doc_id):
    with open(Paths.DATA_DIR / "catalogue-metadata.json") as f:
        cmeta = json.load(f)
    with open(Paths.DATA_DIR / "files-metadata.json") as f:
        fmeta = json.load(f)

    for fm in fmeta:
        if doc_id == fm["id"]:
            parent_id = fm["parent_id"]

    for cm in cmeta:
        if parent_id == cm["id"]:
            return {"title": cm["title"], "notes": cm["notes"]}


def main():
    Paths.NOTES_DIR.mkdir(exist_ok=True, parents=True)

    for file in Paths.PDF_DIR.iterdir():
        out_file = Paths.NOTES_DIR / file.with_suffix(".txt").name
        if out_file.exists():
            continue
        text = extract_text_from_document(file)
        with open(out_file, "w") as f:
            f.write(text)


if __name__ == "__main__":
    main()
